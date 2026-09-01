from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "content" / "partner" / "partneravtal-v1.md"
OUTPUT = ROOT / "content" / "partner" / "Handymate-Partneravtal-v1.docx"
LOGO = ROOT / "public" / "marketing" / "content-library-v1" / "brand" / "handymate-mark.png"

# contract_negotiation_brief preset, with named Swedish A4 + Handymate brand override.
TEAL = "0F7F79"
TEAL_DARK = "075E5A"
TEAL_LIGHT = "E8F5F3"
NAVY = "0B1634"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F4F6F9"
BORDER = "D8E1E7"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sida ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def set_run_font(run, size=11, color=INK, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_markdown(paragraph, text: str, *, size=11, color=INK, italic=False) -> None:
    chunks = re.split(r"(\*\*[^*]+\*\*)", text)
    for chunk in chunks:
        if not chunk:
            continue
        is_bold = chunk.startswith("**") and chunk.endswith("**")
        value = chunk[2:-2] if is_bold else chunk
        run = paragraph.add_run(value)
        set_run_font(run, size=size, color=color, bold=is_bold, italic=italic)


def set_paragraph_border_bottom(paragraph, color=BORDER, size=8, space=5) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Named A4 override: 6.5-inch content width, centered horizontally.
    section.left_margin = Inches(0.8835)
    section.right_margin = Inches(0.8835)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, TEAL_DARK, 14, 8),
        "Heading 2": (13, TEAL_DARK, 11, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True
        style.paragraph_format.left_indent = Pt(0)
        style.paragraph_format.right_indent = Pt(0)
        style.paragraph_format.first_line_indent = Pt(0)

    doc.core_properties.title = "Handymate Partneravtal"
    doc.core_properties.subject = "Hänvisnings- och marknadsföringspartner"
    doc.core_properties.author = "Handymate"
    doc.core_properties.keywords = "partneravtal, provision, Handymate"
    doc.core_properties.comments = "Version 1.0 — för juridisk slutgranskning före publicering"


def configure_running_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("HANDYMATE")
    set_run_font(r1, size=8.5, color=TEAL_DARK, bold=True)
    r2 = p.add_run("   |   PARTNERAVTAL V1.0")
    set_run_font(r2, size=8.5, color=MUTED)
    set_paragraph_border_bottom(p, color=BORDER, size=5, space=4)

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(26)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(0.78))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("AVTALSUTKAST · FÖR JURIDISK SLUTGRANSKNING")
    set_run_font(r, size=9, color=TEAL_DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Handymate\nPartneravtal")
    set_run_font(r, size=31, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Hänvisnings- och marknadsföringspartners")
    set_run_font(r, size=14, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(22)
    set_paragraph_border_bottom(p, color=TEAL, size=16, space=6)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_widths(table, [1700, 7660])
    rows = [
        ("Version", "1.0 · 1 september 2026"),
        ("Standardmodell", "20 % av nettoabonnemangsintäkt i 36 kalendermånader"),
        ("Samarbetsform", "Icke-exklusiv hänvisnings- och marknadsföringspartner"),
        ("Dokument", "Allmänna villkor + kommersiella villkor + partnerbekräftelse"),
    ]
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        set_cell_shading(left, TEAL_LIGHT)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        left.text = ""
        right.text = ""
        p1 = left.paragraphs[0]
        add_inline_markdown(p1, f"**{label}**", size=9.5, color=TEAL_DARK)
        p1.paragraph_format.space_after = Pt(0)
        p2 = right.paragraphs[0]
        add_inline_markdown(p2, value, size=9.5, color=INK)
        p2.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("VIKTIGT")
    set_run_font(r, size=9, color=TEAL_DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline_markdown(
        p,
        "Dokumentet är komplett som avtalsunderlag men ska juridiskt slutgranskas och fyllas med korrekt avtalspart innan extern publicering eller acceptans.",
        size=10.5,
        color=INK,
    )

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    p.paragraph_format.page_break_before = text.startswith("Bilaga ") or text.startswith("C. Avtalsuppgifter")
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    add_inline_markdown(p, text, size={1: 16, 2: 13, 3: 12}[min(level, 3)], color=TEAL_DARK)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    add_inline_markdown(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    marker = p.add_run("•\t")
    set_run_font(marker, color=TEAL_DARK, bold=True)
    add_inline_markdown(p, text)


def add_numbered_intro(doc: Document, number: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.375)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    marker = p.add_run(f"{number}\t")
    set_run_font(marker, color=TEAL_DARK, bold=True)
    add_inline_markdown(p, text)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    if col_count == 2:
        widths = [2700, 6660]
    else:
        widths = [CONTENT_WIDTH_DXA // col_count] * col_count
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_widths(table, widths)

    for i, source_row in enumerate(rows):
        for j, value in enumerate(source_row):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = ""
            if i == 0:
                set_cell_shading(cell, TEAL_DARK)
            elif i % 2 == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(
                p,
                value,
                size=9.2,
                color=WHITE if i == 0 else INK,
            )
            if i == 0:
                for run in p.runs:
                    run.bold = True
        if i == 0:
            set_repeat_table_header(table.rows[i])

    if rows[0] == ["För Handymate", "För Partnern"]:
        table.rows[-1].height = Cm(1.55)
        table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def parse_markdown_into_doc(doc: Document, text: str) -> None:
    lines = text.splitlines()
    # Cover already contains title and metadata. Start at the first contract paragraph.
    start = next(i for i, line in enumerate(lines) if line.startswith("Detta partneravtal"))
    i = start
    paragraph_parts: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_parts
        if paragraph_parts:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_parts))
            paragraph_parts = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped == "---":
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
                if not re.fullmatch(r"\|?[\s:|-]+\|?", row)
            ]
            add_table(doc, parsed)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            title = heading.group(2)
            if title == "Handymate Partneravtal":
                i += 1
                continue
            level = len(heading.group(1))
            # Markdown H1 sections are appendices; ordinary H2 is the contract's H1.
            effective = 1 if level <= 2 else 2
            add_heading(doc, title, effective)
            i += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            add_bullet(doc, bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            add_numbered_intro(doc, numbered.group(1) + ".", numbered.group(2))
            i += 1
            continue

        if stripped.startswith("_") and set(stripped) == {"_"}:
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(16)
            p.paragraph_format.line_spacing = 1.25
            p.add_run(" ")
            set_paragraph_border_bottom(p, color=MUTED, size=4, space=1)
            i += 1
            continue

        paragraph_parts.append(stripped)
        i += 1

    flush_paragraph()


def audit_document(doc: Document) -> None:
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.27
    assert round(section.page_height.inches, 2) == 11.69
    assert round(section.left_margin.inches + section.right_margin.inches + 6.5, 2) == 8.27
    assert doc.styles["Normal"].font.name == "Calibri"
    assert doc.styles["Normal"].font.size.pt == 11
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == CONTENT_WIDTH_DXA
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == TABLE_INDENT_DXA


def main() -> None:
    source_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    configure_running_header_footer(doc)
    add_cover(doc)
    parse_markdown_into_doc(doc, source_text)
    audit_document(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
