from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "content" / "partner" / "partneravtal-v1.md"
OUTPUT = ROOT / "content" / "partner" / "Handymate-Partneravtal-v1-Google-Docs.docx"

BLACK = "000000"
DARK_GRAY = "434343"
MID_GRAY = "666666"
LIGHT_GRAY = "D9D9D9"
CONTENT_WIDTH_DXA = 9360


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, *, size=11, color=BLACK, bold=None, italic=None) -> None:
    run.font.name = "Arial"
    r_fonts = run._element.get_or_add_rPr().rFonts
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_markdown(paragraph, text: str, *, size=11, color=BLACK) -> None:
    for chunk in re.split(r"(\*\*[^*]+\*\*)", text):
        if not chunk:
            continue
        is_bold = chunk.startswith("**") and chunk.endswith("**")
        value = chunk[2:-2] if is_bold else chunk
        run = paragraph.add_run(value)
        set_font(run, size=size, color=color, bold=is_bold)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), LIGHT_GRAY)


def set_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    specs = {
        "Heading 1": (20, BLACK, 20, 6),
        "Heading 2": (16, BLACK, 18, 6),
        "Heading 3": (14, DARK_GRAY, 16, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    doc.core_properties.title = "Handymate Partneravtal"
    doc.core_properties.subject = "Hänvisnings- och marknadsföringspartners"
    doc.core_properties.author = "Handymate"
    doc.core_properties.comments = "Google Docs-kompatibel importversion"


def add_document_title(doc: Document) -> None:
    # Deliberately a plain paragraph, not Word's Title style, for Google Docs import.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Handymate Partneravtal")
    set_font(run, size=26, color=BLACK, bold=False)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Hänvisnings- och marknadsföringspartners")
    set_font(run, size=13, color=DARK_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_inline_markdown(p, "**Version 1.0 · 1 september 2026**", size=10, color=MID_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("Avtalsutkast för juridisk slutgranskning före publicering.")
    set_font(run, size=10, color=MID_GRAY, italic=True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.widow_control = True
    p.paragraph_format.keep_with_next = text.rstrip().endswith(":")
    # Keep the long exclusion list visually intact in Word/Google Docs imports.
    p.paragraph_format.page_break_before = text.startswith("3.6 ")
    add_inline_markdown(p, text)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    p.paragraph_format.page_break_before = text.startswith("Bilaga ")
    add_inline_markdown(
        p,
        text,
        size={1: 20, 2: 16, 3: 14}[min(level, 3)],
        color=BLACK if level < 3 else DARK_GRAY,
    )


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_inline_markdown(p, text)
    for run in p.runs:
        set_font(run, size=11)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_inline_markdown(p, text)
    for run in p.runs:
        set_font(run, size=11)


def add_signature_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("________________________________________")
    set_font(run, size=11, color=MID_GRAY)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [CONTENT_WIDTH_DXA // col_count] * col_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_widths(table, widths)
    set_table_borders(table)
    for row_idx, source_row in enumerate(rows):
        for col_idx, value in enumerate(source_row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(p, value, size=10)
            if row_idx == 0:
                for run in p.runs:
                    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def parse_markdown(doc: Document, source: str) -> None:
    lines = source.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("Detta partneravtal"))
    i = start
    paragraph_parts: list[str] = []

    def flush() -> None:
        nonlocal paragraph_parts
        if paragraph_parts:
            add_body(doc, " ".join(part.strip() for part in paragraph_parts))
            paragraph_parts = []

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            flush()
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush()
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for row in table_lines:
                cells = [cell.strip() for cell in row.strip("|").split("|")]
                if all(re.fullmatch(r"[: -]+", cell) for cell in cells):
                    continue
                rows.append(cells)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush()
            title = heading.group(2)
            if title != "Handymate Partneravtal":
                source_level = len(heading.group(1))
                effective_level = 1 if source_level <= 2 else 2
                add_heading(doc, title, effective_level)
            i += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            flush()
            add_bullet(doc, bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            flush()
            add_numbered(doc, numbered.group(1))
            i += 1
            continue
        if stripped.startswith("_") and set(stripped) == {"_"}:
            flush()
            add_signature_line(doc)
            i += 1
            continue
        paragraph_parts.append(stripped.replace("  ", " "))
        i += 1
    flush()


def audit(doc: Document) -> None:
    assert doc.styles["Normal"].font.name == "Arial"
    assert doc.styles["Normal"].font.size.pt == 11
    assert doc.paragraphs[0].style.name == "Normal"
    assert doc.paragraphs[0].text == "Handymate Partneravtal"
    assert doc.paragraphs[0].runs[0].font.size.pt == 26
    assert not any(p.text.strip() for section in doc.sections for p in section.header.paragraphs)
    assert not any(p.text.strip() for section in doc.sections for p in section.footer.paragraphs)
    for table in doc.tables:
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == "0"


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_document_title(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    audit(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
